"""
Document parsing utilities.
Auto-detects file type and routes to the correct parser, returning
plain text that can be chunked and embedded downstream.
"""

import os
from pypdf import PdfReader
from docx import Document
import pandas as pd


class UnsupportedFileTypeError(Exception):
    pass


class EmptyDocumentError(Exception):
    pass


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt"}


def get_file_extension(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def parse_pdf(file_bytes_or_path) -> str:
    reader = PdfReader(file_bytes_or_path)
    text_parts = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            text_parts.append(f"[Page {i + 1}]\n{page_text}")
    return "\n\n".join(text_parts)


def parse_docx(file_bytes_or_path) -> str:
    doc = Document(file_bytes_or_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of tables, since python-docx skips them by default
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n".join(parts)


def parse_excel_or_csv(file_bytes_or_path, extension: str) -> str:
    if extension == ".csv":
        df = pd.read_csv(file_bytes_or_path)
        return df.to_string(index=False)

    # Excel may have multiple sheets — read them all
    sheets = pd.read_excel(file_bytes_or_path, sheet_name=None)
    parts = []
    for sheet_name, df in sheets.items():
        parts.append(f"[Sheet: {sheet_name}]\n{df.to_string(index=False)}")
    return "\n\n".join(parts)


def parse_txt(file_bytes_or_path) -> str:
    if hasattr(file_bytes_or_path, "read"):
        raw = file_bytes_or_path.read()
        if isinstance(raw, bytes):
            return raw.decode("utf-8", errors="ignore")
        return raw
    with open(file_bytes_or_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def parse_document(uploaded_file) -> dict:
    """
    Takes a Streamlit UploadedFile (or file path string) and returns:
        {"filename": str, "text": str, "num_chars": int}
    Raises UnsupportedFileTypeError / EmptyDocumentError on bad input.
    """
    if isinstance(uploaded_file, str):
        filename = os.path.basename(uploaded_file)
        source = uploaded_file
    else:
        filename = uploaded_file.name
        source = uploaded_file

    ext = get_file_extension(filename)

    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"'{ext}' is not supported. Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    try:
        if ext == ".pdf":
            text = parse_pdf(source)
        elif ext == ".docx":
            text = parse_docx(source)
        elif ext in (".xlsx", ".xls", ".csv"):
            text = parse_excel_or_csv(source, ext)
        elif ext == ".txt":
            text = parse_txt(source)
        else:
            raise UnsupportedFileTypeError(f"No parser wired up for '{ext}'")
    except UnsupportedFileTypeError:
        raise
    except Exception as e:
        raise ValueError(f"Failed to parse '{filename}': {e}")

    if not text or not text.strip():
        raise EmptyDocumentError(f"'{filename}' appears to be empty or has no extractable text.")

    return {"filename": filename, "text": text, "num_chars": len(text)}
